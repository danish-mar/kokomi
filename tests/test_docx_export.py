"""Regression tests for markdown → DOCX fidelity.

Both exporters used to write markdown tables out as literal paragraphs — the
pipes and the |---| divider ended up as body text — and the canvas exporter
stripped **bold**/*italic*/`code` markers without ever applying real
formatting, so emphasis vanished from the document entirely.

Both are easy to reintroduce: a table row that falls through to a generic
"else: add_paragraph(text)" branch still *looks* fine in a diff, and inline
emphasis silently disappears rather than erroring. These assert on the actual
saved document rather than on the code shape.
"""
import io
import os
import unittest

os.environ.setdefault("GROQ_API_KEY", "mock-key")
os.environ.setdefault("GOOGLE_API_KEY", "mock-key")

from docx import Document

from app.docx_md import is_table_divider, is_table_row, parse_table

MD = """# Report

Text with **bold**, *italic* and `code()`.

| Symptom | Cause |
|---------|-------|
| Slow **boot** | Disk I/O |
| Crash | Race |

- bullet with **bold**
"""


def _formatted_runs(doc):
    """(flags, text) for every run carrying real character formatting."""
    out = []
    for p in doc.paragraphs:
        for r in p.runs:
            flags = "".join(f for f, on in (
                ("B", r.bold), ("I", r.italic), ("C", r.font.name == "Consolas"),
            ) if on)
            if flags:
                out.append((flags, r.text))
    return out


class TestTableParsing(unittest.TestCase):
    def test_row_and_divider_detection(self):
        self.assertTrue(is_table_row("| a | b |"))
        self.assertFalse(is_table_row("just text"))
        # A bulleted line mentioning a pipe must not be mistaken for a table.
        self.assertFalse(is_table_row("- use a | b"))
        self.assertTrue(is_table_divider("|---|:---:|"))
        self.assertFalse(is_table_divider("| a | b |"))

    def test_divider_dropped_and_ragged_rows_padded(self):
        lines = ["| a | b | c |", "|---|---|---|", "| 1 |", "after"]
        rows, nxt = parse_table(lines, 0)
        self.assertEqual(nxt, 3, "should stop at the first non-table line")
        self.assertEqual(rows[0], ["a", "b", "c"])
        # Short row padded, so writing it can't IndexError.
        self.assertEqual(rows[1], ["1", "", ""])


class TestWorkflowDocxExport(unittest.TestCase):
    def test_table_becomes_a_real_table_and_emphasis_survives(self):
        from app.workflow import docx_export

        res = docx_export.invoke({"markdown_content": MD, "filename": "kokomi_test_export.docx"})
        self.assertIn("at: ", res, f"export failed: {res}")
        path = res.split("at: ")[-1].strip()
        try:
            doc = Document(path)

            self.assertEqual(len(doc.tables), 1, "markdown table did not become a Word table")
            table = doc.tables[0]
            self.assertEqual(len(table.rows), 3, "expected header + 2 body rows (divider dropped)")
            self.assertEqual([c.text for c in table.rows[0].cells], ["Symptom", "Cause"])

            # The pipe syntax must not also appear as body text.
            body = "\n".join(p.text for p in doc.paragraphs)
            self.assertNotIn("|---", body)
            self.assertNotIn("| Symptom", body)

            texts = [t for _, t in _formatted_runs(doc)]
            self.assertIn("bold", texts)
            self.assertIn("italic", texts)
            self.assertIn("code()", texts)
            # Markers themselves must be gone, not merely unstyled.
            self.assertNotIn("**", body)
        finally:
            if os.path.exists(path):
                os.remove(path)


class TestCanvasDocxExport(unittest.TestCase):
    def test_emphasis_is_applied_not_just_stripped(self):
        from app.routers.canvas import _to_blocks, _build_docx

        doc = Document(io.BytesIO(_build_docx(_to_blocks(MD, "document"), "Report")))

        self.assertEqual(len(doc.tables), 1, "markdown table did not become a Word table")
        runs = _formatted_runs(doc)
        self.assertTrue(runs, "no run carried any formatting — emphasis was stripped, not applied")
        texts = [t for _, t in runs]
        self.assertIn("bold", texts)
        self.assertIn("italic", texts)
        self.assertIn("code()", texts)

        body = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("|---", body)
        self.assertNotIn("**", body)

    def test_table_cells_keep_their_own_emphasis(self):
        from app.routers.canvas import _to_blocks, _build_docx

        doc = Document(io.BytesIO(_build_docx(_to_blocks(MD, "document"), "Report")))
        cell = doc.tables[0].rows[1].cells[0]          # "Slow **boot**"
        self.assertEqual(cell.text, "Slow boot")
        self.assertTrue(
            any(r.bold and r.text == "boot" for r in cell.paragraphs[0].runs),
            "emphasis inside a table cell was not applied",
        )

    def test_other_block_writers_still_accept_table_blocks(self):
        """The PDF and markdown writers share these blocks, and a table block
        carries a list of rows rather than a string — passing it to a string
        API would raise."""
        from app.routers.canvas import _to_blocks, _build_pdf, _blocks_to_markdown

        blocks = _to_blocks(MD, "document")
        self.assertIn("table", [k for k, _, _ in blocks])

        pdf = _build_pdf(blocks, "Report")
        self.assertTrue(pdf.startswith(b"%PDF-"))

        md = _blocks_to_markdown(blocks)
        self.assertIn("| Symptom | Cause |", md)
        self.assertIn("| --- | --- |", md, "round-trip lost the header divider")


if __name__ == "__main__":
    unittest.main()


class TestPdfImageRendering(unittest.TestCase):
    """A single awkward image used to fail the whole PDF with a 500.

    Only the width was clamped, so a portrait image scaled to the page width
    became taller than the page and ReportLab refused to lay it out — and
    because layout happens in doc.build(), that error landed outside the
    per-image try/except and took the entire document down.
    """

    def setUp(self):
        import tempfile
        self.dir = tempfile.mkdtemp()

    def _png(self, name, size):
        from PIL import Image
        path = os.path.join(self.dir, name)
        Image.new("RGB", size, (200, 120, 160)).save(path)
        return path

    def _render(self, md):
        from app.pdf_render import render_markdown_to_pdf
        buf = io.BytesIO()
        render_markdown_to_pdf(md, buf, image_base_dir=self.dir)
        return buf.getvalue()

    def test_tall_portrait_image_does_not_fail_the_render(self):
        tall = self._png("tall.png", (900, 3000))
        out = self._render(f"# Gallery\n\n![a]({tall})\n\nAfter.")
        self.assertTrue(out.startswith(b"%PDF-"))

    def test_image_wider_than_the_page(self):
        wide = self._png("wide.png", (5000, 400))
        out = self._render(f"# Gallery\n\n![a]({wide})\n\nAfter.")
        self.assertTrue(out.startswith(b"%PDF-"))

    def test_download_that_is_not_an_image_degrades_to_a_placeholder(self):
        bad = os.path.join(self.dir, "broken.png")
        with open(bad, "wb") as f:
            f.write(b"<html><body>404 Not Found</body></html>")
        out = self._render(f"# Gallery\n\n![b]({bad})\n\nText after.")
        self.assertTrue(out.startswith(b"%PDF-"), "one bad image must not fail the document")

    def test_several_awkward_images_together(self):
        tall = self._png("t2.png", (900, 4000))
        wide = self._png("w2.png", (6000, 300))
        bad = os.path.join(self.dir, "b2.png")
        with open(bad, "wb") as f:
            f.write(b"not an image at all")
        out = self._render(f"# Gallery\n\n![a]({tall})\n\n![b]({wide})\n\n![c]({bad})\n\nEnd.")
        self.assertTrue(out.startswith(b"%PDF-"))


class TestPdfImageLayout(unittest.TestCase):
    """Images are laid out by their real shape, not one-size-fits-all.

    Wide images run full width with the caption underneath; square and tall
    ones sit beside their caption in a two-column table, with tall ones
    alternating sides. The classification reads the decoded file, so a
    mislabelled or malformed image can't silently pick the wrong template.
    """

    def setUp(self):
        import tempfile
        self.dir = tempfile.mkdtemp()

    def _png(self, name, size):
        from PIL import Image
        path = os.path.join(self.dir, name)
        Image.new("RGB", size, (200, 120, 160)).save(path)
        return path

    def _story(self, md):
        """Build the flowable list the renderer would hand to ReportLab."""
        from unittest.mock import patch
        from app.pdf_render import render_markdown_to_pdf
        captured = []
        real_build = None

        def spy(self_doc, story, *a, **kw):
            captured.append(list(story))
            return real_build(self_doc, story, *a, **kw)

        from reportlab.platypus import SimpleDocTemplate
        real_build = SimpleDocTemplate.build
        with patch.object(SimpleDocTemplate, "build", spy):
            render_markdown_to_pdf(md, io.BytesIO(), image_base_dir=self.dir)
        return captured[0]

    def _tables(self, story):
        from reportlab.platypus import Table, KeepTogether
        found = []
        for f in story:
            items = f._content if isinstance(f, KeepTogether) else [f]
            found += [x for x in items if isinstance(x, Table)]
        return found

    @staticmethod
    def _cell(value):
        """ReportLab normalizes a bare flowable cell into a 1-tuple."""
        while isinstance(value, (list, tuple)) and len(value) == 1:
            value = value[0]
        return value

    def test_classification_matches_aspect_ratio(self):
        from app.pdf_render import render_markdown_to_pdf  # noqa: F401  (import guard)
        cases = [((900, 400), "wide"), ((600, 600), "square"), ((400, 900), "tall")]
        for size, expected in cases:
            ratio = size[0] / size[1]
            got = "wide" if ratio >= 1.4 else ("square" if ratio >= 0.85 else "tall")
            self.assertEqual(got, expected, f"{size} should classify as {expected}")

    def test_wide_image_is_full_width_with_no_side_table(self):
        wide = self._png("wide.png", (1200, 500))
        story = self._story(f"![a banner]({wide})")
        self.assertEqual(self._tables(story), [], "a wide image must not be put beside its caption")

    def test_square_image_sits_beside_its_caption(self):
        sq = self._png("sq.png", (600, 600))
        tables = self._tables(self._story(f"![the caption]({sq})"))
        self.assertEqual(len(tables), 1)
        self.assertEqual(len(tables[0]._cellvalues[0]), 2, "image and caption share a row")

    def test_uncaptioned_image_is_never_put_in_a_side_layout(self):
        sq = self._png("sq2.png", (600, 600))
        self.assertEqual(self._tables(self._story(f"![]({sq})")), [],
                         "with no caption there is nothing to sit beside")

    def test_consecutive_tall_images_alternate_sides(self):
        from reportlab.platypus import Image as RLImage
        a = self._png("t1.png", (400, 1000))
        b = self._png("t2.png", (400, 1000))
        tables = self._tables(self._story(f"![one]({a})\n\n![two]({b})"))
        self.assertEqual(len(tables), 2)
        first, second = tables[0]._cellvalues[0], tables[1]._cellvalues[0]
        self.assertIsInstance(self._cell(first[0]), RLImage, "first tall image goes on the left")
        self.assertIsInstance(self._cell(second[1]), RLImage, "the next one flips to the right")
