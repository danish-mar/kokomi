"""Shared Markdown → DOCX primitives.

Used by both the multi-agent workflow's `docx_export` tool (app/workflow.py)
and the canvas document export (app/routers/canvas.py), so there is exactly
one implementation of "turn markdown into Word content" rather than two copies
drifting apart — the same reasoning as app/pdf_render.py.

Covers the two things plain paragraph writing loses: inline emphasis (which
has to become styled *runs*, since Word has no inline markup) and tables
(which have to become a real table object, not a line of text with pipes in
it).
"""
import re

# Header fill for generated tables, matching the PDF renderer's palette.
TABLE_HEADER_BG = "505081"
TABLE_HEADER_FG = (0xFF, 0xFF, 0xFF)

_INLINE_RE = re.compile(
    r"(\*\*\*.+?\*\*\*|___.+?___"      # bold+italic
    r"|\*\*.+?\*\*|__.+?__"            # bold
    r"|\*[^*]+?\*|(?<!\w)_[^_]+?_(?!\w)"  # italic
    r"|`[^`]+?`"                       # inline code
    r"|\[[^\]]+?\]\([^)]*?\))",        # link
    re.DOTALL,
)
_LINK_RE = re.compile(r"^\[([^\]]+)\]\(([^)]*)\)$")


def add_inline_runs(paragraph, text, size=None, color=None, force_bold=False,
                    font_name="Arial"):
    """Append `text` to `paragraph` as runs carrying real Word formatting.

    Markdown emphasis can't survive as characters — Word needs one run per
    styled span — so writing the stripped text into a single run (or leaving
    the asterisks in) is what makes formatting appear to "not persist".
    """
    from docx.shared import Pt, RGBColor  # noqa: F401  (RGBColor used by callers)

    def _style(run, *, bold=False, italic=False, mono=False, link=False):
        run.bold = bold or force_bold
        run.italic = italic
        run.font.name = "Consolas" if mono else font_name
        if size is not None:
            run.font.size = Pt(size - 0.5 if mono else size)
        if link:
            from docx.shared import RGBColor as _RGB
            run.font.color.rgb = _RGB(0x50, 0x50, 0x81)
            run.underline = True
        elif color is not None:
            run.font.color.rgb = color

    for part in _INLINE_RE.split(text or ""):
        if not part:
            continue
        if (part.startswith("***") and part.endswith("***")) or \
           (part.startswith("___") and part.endswith("___")):
            _style(paragraph.add_run(part[3:-3]), bold=True, italic=True)
        elif (part.startswith("**") and part.endswith("**")) or \
             (part.startswith("__") and part.endswith("__")):
            _style(paragraph.add_run(part[2:-2]), bold=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            _style(paragraph.add_run(part[1:-1]), mono=True)
        elif _LINK_RE.match(part):
            _style(paragraph.add_run(_LINK_RE.match(part).group(1)), link=True)
        elif (part.startswith("*") and part.endswith("*") and len(part) > 2) or \
             (part.startswith("_") and part.endswith("_") and len(part) > 2):
            _style(paragraph.add_run(part[1:-1]), italic=True)
        else:
            _style(paragraph.add_run(part))


def is_table_row(line: str) -> bool:
    s = (line or "").strip()
    return s.startswith("|") and s.count("|") >= 2


def is_table_divider(line: str) -> bool:
    """The |---|:--:| row under a table's header."""
    s = (line or "").strip()
    if not is_table_row(s):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return bool(cells) and all(c and set(c) <= set("-: ") for c in cells)


def parse_table(lines, start: int):
    """Consume a run of table rows beginning at `lines[start]`.

    Returns (rows, next_index). `rows` is a list of cell-string lists with the
    divider row dropped; every row is padded to the widest one so a ragged
    table can't raise IndexError while being written.
    """
    rows, i = [], start
    while i < len(lines) and is_table_row(lines[i]):
        if not is_table_divider(lines[i]):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    if rows:
        width = max(len(r) for r in rows)
        for r in rows:
            r.extend([""] * (width - len(r)))
    return rows, i


def _shade(cell, hex_fill: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, rows, size=9.5, header=True):
    """Write `rows` into `doc` as a real Word table with a shaded header."""
    from docx.shared import RGBColor

    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Table Grid"   # gives every cell a visible border
    except KeyError:
        pass                          # template without that style; borders omitted

    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            # A new cell already holds one empty paragraph; write into it
            # instead of adding a second, which would double the row height.
            para = cell.paragraphs[0]
            is_head = header and r == 0
            add_inline_runs(
                para, text, size=size,
                color=RGBColor(*TABLE_HEADER_FG) if is_head else None,
                force_bold=is_head,
            )
            if is_head:
                _shade(cell, TABLE_HEADER_BG)
    return table
