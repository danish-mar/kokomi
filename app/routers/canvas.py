"""Canvas persistence.

A canvas is an artifact the user can edit in place (a code editor or a
Word-style document pane) rather than just read. Edits are written back onto
the artifact stored on the assistant message that produced it, which is also
what gets injected into the system prompt on the next turn — so the model
always sees the user's current version, not the one it originally wrote.
"""

import csv as csv_module
import io
import json
import re
from html import escape as html_escape
from html.parser import HTMLParser
from typing import Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

from app.storage import load_convos, save_convos

router = APIRouter(prefix="/api")


# ── Content → structured blocks ───────────────────────────────────────
# A document canvas holds Quill HTML once edited, but plain markdown when it
# arrives fresh from the model. Both are normalised to the same block list so
# the DOCX and PDF writers only deal with one shape.
#
# A block is (kind, text, level) where kind is one of:
#   heading | para | bullet | number | quote | code


class _QuillHtmlParser(HTMLParser):
    """Flatten Quill's HTML into blocks. Quill emits a shallow, predictable
    tree (h1-h6 / p / ul>li / ol>li / blockquote / pre), so a streaming parser
    is enough — no need for a full DOM library."""

    def __init__(self):
        super().__init__()
        self.blocks: list[tuple[str, str, int]] = []
        self._kind = "para"
        self._level = 0
        self._buf: list[str] = []
        self._list_stack: list[str] = []

    def _flush(self):
        text = re.sub(r"[ \t]+", " ", "".join(self._buf)).strip()
        if text:
            self.blocks.append((self._kind, text, self._level))
        self._buf = []
        self._kind = "para"
        self._level = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("ul", "ol"):
            self._list_stack.append(tag)
        elif tag == "li":
            self._flush()
            self._kind = "number" if (self._list_stack and self._list_stack[-1] == "ol") else "bullet"
        elif re.fullmatch(r"h[1-6]", tag):
            self._flush()
            self._kind, self._level = "heading", int(tag[1])
        elif tag == "blockquote":
            self._flush()
            self._kind = "quote"
        elif tag == "pre":
            self._flush()
            self._kind = "code"
        elif tag == "p":
            self._flush()
        elif tag == "br":
            self._buf.append("\n")

    def handle_endtag(self, tag):
        if tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
        elif tag in ("li", "p", "blockquote", "pre") or re.fullmatch(r"h[1-6]", tag):
            self._flush()

    def handle_data(self, data):
        self._buf.append(data)

    def close(self):
        super().close()
        self._flush()


def _markdown_to_blocks(md: str) -> list[tuple[str, str, int]]:
    blocks: list[tuple[str, str, int]] = []
    in_code = False
    for raw in (md or "").split("\n"):
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            blocks.append(("code", raw, 0))
            continue
        s = line.strip()
        if not s:
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", s)
        if h:
            blocks.append(("heading", h.group(2), len(h.group(1))))
        elif re.match(r"^[-*+]\s+", s):
            blocks.append(("bullet", re.sub(r"^[-*+]\s+", "", s), 0))
        elif re.match(r"^\d+[.)]\s+", s):
            blocks.append(("number", re.sub(r"^\d+[.)]\s+", "", s), 0))
        elif s.startswith(">"):
            blocks.append(("quote", s.lstrip("> ").strip(), 0))
        else:
            blocks.append(("para", s, 0))
    return blocks


def _looks_like_html(text: str) -> bool:
    return bool(re.match(r"\s*<(p|h[1-6]|ul|ol|blockquote|div|pre|table)[\s>]", text or "", re.I))


def _to_blocks(content: str, mode: str) -> list[tuple[str, str, int]]:
    if mode == "code":
        return [("code", line, 0) for line in (content or "").split("\n")]
    if _looks_like_html(content):
        p = _QuillHtmlParser()
        p.feed(content or "")
        p.close()
        return p.blocks
    return _markdown_to_blocks(content)


def _strip_inline_md(text: str) -> str:
    """Drop **bold**/*italic*/`code` markers — the writers apply real styling."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    return re.sub(r"`(.+?)`", r"\1", text)


class CanvasUpdate(BaseModel):
    content: str
    title: Optional[str] = None


class CanvasEdit(BaseModel):
    """A targeted edit: rewrite `selection` per `instruction`, nothing else."""
    instruction: str
    selection: Optional[str] = ""


class CanvasPatch(BaseModel):
    """A line- (code), block- (document) or range- (spreadsheet) addressed
    patch request over the canvas."""
    instruction: str
    # Optional line window the user highlighted, to focus the edit.
    start_line: Optional[int] = None
    end_line: Optional[int] = None
    # Optional A1-style range the user had selected in a spreadsheet canvas.
    range: Optional[str] = None


def _locate(convos: dict, conversation_id: str, artifact_id: str):
    """Return (conversation, artifact) for an artifact id, newest version first.

    A canvas is re-emitted under the same id whenever the model updates it, so
    later messages can hold newer copies — scan back to front.
    """
    convo = convos.get(conversation_id)
    if not convo:
        raise HTTPException(status_code=404, detail="Conversation not found")

    for msg in reversed(convo.get("messages") or []):
        for art in (msg.get("artifacts") or []):
            if art.get("id") == artifact_id:
                return convo, art

    raise HTTPException(status_code=404, detail="Canvas not found")


@router.get("/canvas/{conversation_id}/{artifact_id}")
async def get_canvas(conversation_id: str, artifact_id: str):
    _, art = _locate(load_convos(), conversation_id, artifact_id)
    return {
        "id": art.get("id"),
        "title": art.get("title"),
        "mode": art.get("mode") or "code",
        "language": art.get("language") or "plaintext",
        "content": art.get("content") or "",
    }


@router.put("/canvas/{conversation_id}/{artifact_id}")
async def update_canvas(conversation_id: str, artifact_id: str, body: CanvasUpdate):
    """Save user edits back onto the stored artifact."""
    convos = load_convos()
    _, art = _locate(convos, conversation_id, artifact_id)

    art["content"] = body.content
    if body.title:
        art["title"] = body.title
    art["edited_by_user"] = True

    save_convos(convos)
    return {"ok": True, "id": artifact_id, "bytes": len(body.content)}


# ── Targeted edit ─────────────────────────────────────────────────────

def _strip_wrapper(text: str) -> str:
    """Models like to gift-wrap replacements. Unwrap fences and stray quotes."""
    t = (text or "").strip()
    fence = re.match(r"^```[\w-]*\n(.*)\n```$", t, re.S)
    if fence:
        t = fence.group(1)
    # Only unquote when the WHOLE reply is quoted, not a quote inside prose.
    if len(t) > 1 and t[0] == t[-1] and t[0] in "\"'" and t.count(t[0]) == 2:
        t = t[1:-1]
    return t.strip()


@router.post("/canvas/{conversation_id}/{artifact_id}/edit")
async def edit_canvas(conversation_id: str, artifact_id: str, body: CanvasEdit):
    """Rewrite just the selected passage — no chat turn, no whole-document rewrite.

    Returns only replacement text, which the editor splices in place. The full
    canvas is passed as context so the rewrite matches surrounding voice, but
    the model is told in no uncertain terms to return the passage alone.
    """
    from langchain_core.messages import SystemMessage, HumanMessage
    from app.llm import get_llm
    from app.storage import load_prefs

    _, art = _locate(load_convos(), conversation_id, artifact_id)
    mode = (art.get("mode") or "code").lower()
    selection = (body.selection or "").strip()
    instruction = (body.instruction or "").strip()
    if not instruction:
        raise HTTPException(status_code=400, detail="instruction is required")

    # Context is for tone/consistency only; keep it bounded.
    full = art.get("content") or ""
    if mode == "document" and _looks_like_html(full):
        full = "\n".join(b[1] for b in _to_blocks(full, mode))
    context = full[:6000]

    if selection:
        system = (
            "You are a precise document editor. You are given a passage from a larger "
            "document and an instruction. Rewrite ONLY that passage.\n"
            "Output rules — these are absolute:\n"
            "- Return ONLY the rewritten passage, nothing else.\n"
            "- No preamble, no explanation, no commentary, no code fences, no surrounding quotes.\n"
            "- Do NOT return the whole document. Do NOT repeat unrelated parts.\n"
            "- Keep the same language and roughly the same formatting style. Plain prose in, "
            "plain prose out; markdown inline formatting (**bold**, *italic*) is allowed.\n"
            "- If the instruction cannot sensibly apply, return the passage unchanged."
        )
        human = (
            f"DOCUMENT (context only — do not return this):\n{context}\n\n"
            f"INSTRUCTION: {instruction}\n\n"
            f"PASSAGE TO REWRITE:\n{selection}"
        )
    else:
        system = (
            "You are a precise writing assistant working inside a document editor. "
            "Produce text to insert at the user's cursor.\n"
            "Output rules — these are absolute:\n"
            "- Return ONLY the new text to insert, nothing else.\n"
            "- No preamble, no explanation, no code fences.\n"
            "- Do NOT repeat the existing document.\n"
            "- Match the document's voice, tense and formatting."
        )
        human = (
            f"DOCUMENT (context — do not repeat it):\n{context}\n\n"
            f"INSTRUCTION: {instruction}"
        )

    try:
        llm = get_llm(load_prefs())
        resp = await llm.ainvoke([SystemMessage(content=system), HumanMessage(content=human)])
        text = _strip_wrapper(getattr(resp, "content", "") or "")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Edit failed: {e}")

    if not text:
        raise HTTPException(status_code=502, detail="Model returned an empty edit")

    return {"text": text, "replaced": bool(selection)}


# ── Line-addressed patching ───────────────────────────────────────────
#
# Instead of regenerating the file, the model is shown numbered lines and asked
# for a small set of edits. Each edit carries an `expect` anchor — the first
# line it believes it is replacing — which is verified against the real content
# before anything is applied. That anchor is the whole point: line numbers are
# easy for a model to be off-by-one on, and applying a confidently-wrong range
# silently corrupts the file. If the anchor doesn't match, the edit is refused
# rather than guessed at.


# Documents are addressed by BLOCK, not by line. Quill stores its content as a
# flat run of top-level elements with no newlines at all, so the whole document
# is "line 1" and a line-addressed edit would replace everything — exactly the
# full rewrite this feature exists to avoid. Paragraphs and headings are to
# prose what lines are to code, so those are the addressable unit.

_BLOCK_RE = re.compile(
    r"<(p|h[1-6]|ul|ol|blockquote|pre|div)(?:\s[^>]*)?>(.*?)</\1>", re.S | re.I)


def split_html_blocks(html: str) -> list:
    """Flat list of (tag, inner_html, plain_text) for each top-level block."""
    blocks = []
    for m in _BLOCK_RE.finditer(html or ""):
        tag = m.group(1).lower()
        inner = m.group(2)
        text = re.sub(r"<[^>]+>", " ", inner)
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            blocks.append({"tag": tag, "inner": inner, "text": text, "span": m.span()})
    return blocks


def _number_blocks(html: str) -> str:
    out = []
    for i, b in enumerate(split_html_blocks(html), 1):
        out.append(f"[{i}] <{b['tag']}> {b['text']}")
    return "\n".join(out)


def _inline_md_to_html(text: str) -> str:
    """Minimal inline markdown so **bold**/*italic*/`code` survive a rewrite."""
    from html import escape
    t = escape(text or "")
    t = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", t)
    t = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", t)
    t = re.sub(r"`(.+?)`", r"<code>\1</code>", t)
    return t


def _render_block(tag: str, replacement: str) -> str:
    """Rebuild a block, keeping its original tag so structure is preserved."""
    if tag in ("ul", "ol"):
        items = [ln.strip().lstrip("-*+ ").strip()
                 for ln in replacement.split("\n") if ln.strip()]
        lis = "".join(f"<li>{_inline_md_to_html(i)}</li>" for i in items)
        return f"<{tag}>{lis}</{tag}>"
    if tag == "pre":
        from html import escape
        return f"<pre>{escape(replacement)}</pre>"
    # Multi-paragraph replacement of a single block: emit sibling <p>s.
    parts = [p.strip() for p in replacement.split("\n\n") if p.strip()]
    if len(parts) > 1 and tag == "p":
        return "".join(f"<p>{_inline_md_to_html(p)}</p>" for p in parts)
    return f"<{tag}>{_inline_md_to_html(replacement)}</{tag}>"


def apply_block_edits(html: str, edits: list) -> tuple:
    """Apply block-addressed edits to Quill HTML. Mirrors apply_line_edits:
    anchors are verified, edits apply bottom-up, bad ones are refused."""
    blocks = split_html_blocks(html)
    total = len(blocks)
    normalized, rejected = [], []

    for e in edits:
        try:
            idx = int(e.get("block"))
        except (TypeError, ValueError):
            rejected.append({"reason": "block is not an integer"})
            continue
        if idx < 1 or idx > total:
            rejected.append({"reason": f"block {idx} out of range (1..{total})"})
            continue

        expect = e.get("expect")
        actual = blocks[idx - 1]["text"]
        if expect and _norm(expect) not in _norm(actual) and _norm(actual) not in _norm(expect):
            rejected.append({
                "reason": f"anchor mismatch at block {idx}: "
                          f"expected {expect[:60]!r}, found {actual[:60]!r}"})
            continue
        normalized.append({"idx": idx, "replacement": e.get("replacement", "") or ""})

    seen = set()
    kept = []
    for ed in normalized:
        if ed["idx"] in seen:
            rejected.append({"reason": f"duplicate edit for block {ed['idx']}"})
            continue
        seen.add(ed["idx"])
        kept.append(ed)

    new_html = html
    applied = []
    for ed in sorted(kept, key=lambda x: x["idx"], reverse=True):
        b = blocks[ed["idx"] - 1]
        start, end = b["span"]
        rendered = "" if ed["replacement"] == "" else _render_block(b["tag"], ed["replacement"])
        new_html = new_html[:start] + rendered + new_html[end:]
        applied.append({"block": ed["idx"], "tag": b["tag"],
                        "deleted": ed["replacement"] == ""})

    applied.reverse()
    return new_html, applied, rejected


def _number_lines(text: str) -> str:
    lines = (text or "").split("\n")
    width = len(str(len(lines)))
    return "\n".join(f"{i:>{width}}\t{ln}" for i, ln in enumerate(lines, 1))


def _norm(s: str) -> str:
    """Compare anchors ignoring whitespace noise the model tends to reflow."""
    return re.sub(r"\s+", " ", (s or "")).strip()


def apply_line_edits(content: str, edits: list) -> tuple[str, list, list]:
    """Apply line-range edits. Returns (new_content, applied, rejected).

    Edits are applied from the bottom up so that earlier line numbers stay
    valid as the document shifts underneath them.
    """
    lines = (content or "").split("\n")
    total = len(lines)

    normalized = []
    rejected = []

    for e in edits:
        try:
            start = int(e.get("start_line"))
            end = int(e.get("end_line", start))
        except (TypeError, ValueError):
            rejected.append({"edit": e, "reason": "start_line/end_line not integers"})
            continue

        replacement = e.get("replacement", "")
        if replacement is None:
            replacement = ""

        is_insert = end == start - 1          # empty range == pure insertion
        if start < 1 or start > total + 1:
            rejected.append({"edit": e, "reason": f"start_line {start} out of range (1..{total})"})
            continue
        if not is_insert and (end < start or end > total):
            rejected.append({"edit": e, "reason": f"end_line {end} out of range ({start}..{total})"})
            continue

        # Verify the anchor before trusting the line numbers.
        expect = e.get("expect")
        if expect and not is_insert:
            actual = lines[start - 1]
            if _norm(expect) not in _norm(actual) and _norm(actual) not in _norm(expect):
                rejected.append({
                    "edit": e,
                    "reason": f"anchor mismatch at line {start}: "
                              f"expected {expect!r}, found {actual!r}",
                })
                continue

        normalized.append({"start": start, "end": end, "replacement": replacement,
                           "is_insert": is_insert})

    # Overlapping ranges would corrupt each other — keep the first, drop the rest.
    normalized.sort(key=lambda x: x["start"])
    kept, last_end = [], 0
    for ed in normalized:
        if not ed["is_insert"] and ed["start"] <= last_end:
            rejected.append({"edit": ed, "reason": "overlaps an earlier edit"})
            continue
        kept.append(ed)
        last_end = max(last_end, ed["end"])

    applied = []
    for ed in sorted(kept, key=lambda x: x["start"], reverse=True):
        new_lines = ed["replacement"].split("\n") if ed["replacement"] != "" else []
        if ed["is_insert"]:
            lines[ed["start"] - 1:ed["start"] - 1] = new_lines
        else:
            lines[ed["start"] - 1:ed["end"]] = new_lines
        applied.append({
            "start_line": ed["start"],
            "end_line": ed["end"],
            "lines_removed": 0 if ed["is_insert"] else (ed["end"] - ed["start"] + 1),
            "lines_added": len(new_lines),
        })

    applied.reverse()
    return "\n".join(lines), applied, rejected


# ── Cell-addressed patching (spreadsheet canvas) ────────────────────────
#
# A spreadsheet canvas stores its content as plain CSV — the same "raw source"
# treatment code gets, rather than reinventing x-spreadsheet's own nested JSON
# model in the prompt. Edits are addressed by A1-style range (e.g. "B2" or
# "B2:D5"), the notation every spreadsheet user and LLM already knows. Unlike
# line edits, writing into cells doesn't shift anything else, so — once
# overlapping ranges are resolved — edits can be applied in any order.

def _parse_grid(text: str) -> list:
    """CSV text -> a rectangular grid (ragged rows padded with '')."""
    rows = [list(r) for r in csv_module.reader(io.StringIO(text or ""))]
    width = max((len(r) for r in rows), default=0)
    for r in rows:
        r.extend([""] * (width - len(r)))
    return rows


def _grid_to_csv(grid: list) -> str:
    buf = io.StringIO()
    csv_module.writer(buf, lineterminator="\n").writerows(grid)
    return buf.getvalue()


def _col_letters(n: int) -> str:
    """0-based column index -> spreadsheet letters (0->A, 25->Z, 26->AA)."""
    s = ""
    n += 1
    while n:
        n, rem = divmod(n - 1, 26)
        s = chr(65 + rem) + s
    return s


def _col_index(letters: str) -> int:
    n = 0
    for ch in letters.upper():
        n = n * 26 + (ord(ch) - 64)
    return n - 1


_A1_RE = re.compile(r"^([A-Za-z]+)(\d+)$")


def _parse_a1(a1: str):
    m = _A1_RE.match((a1 or "").strip())
    if not m:
        return None
    return int(m.group(2)) - 1, _col_index(m.group(1))  # 0-based (row, col)


def _parse_range(range_str: str):
    """'B2' or 'B2:D5' -> (r0, c0, r1, c1) inclusive, 0-based. None if invalid."""
    s = (range_str or "").strip()
    if ":" in s:
        a, b = s.split(":", 1)
        pa, pb = _parse_a1(a), _parse_a1(b)
        if not pa or not pb:
            return None
        return min(pa[0], pb[0]), min(pa[1], pb[1]), max(pa[0], pb[0]), max(pa[1], pb[1])
    p = _parse_a1(s)
    if not p:
        return None
    return p[0], p[1], p[0], p[1]


def _numbered_grid(grid: list, max_rows: int = 200) -> str:
    """Render the grid as a lettered/numbered table, the same 'show
    coordinates, ask for anchored edits' pattern as _number_lines/_number_blocks."""
    width = len(grid[0]) if grid else 0
    header = "     " + " ".join(f"{_col_letters(c):>10}" for c in range(width))
    out = [header]
    for i, row in enumerate(grid[:max_rows]):
        cells = " ".join(f"{(c or ''):>10.10}" for c in row)
        out.append(f"{i + 1:>4} {cells}")
    if len(grid) > max_rows:
        out.append(f"… ({len(grid) - max_rows} more rows not shown)")
    return "\n".join(out)


def apply_cell_edits(content: str, edits: list) -> tuple:
    """Apply A1-range-addressed edits to a CSV grid. Mirrors apply_line_edits/
    apply_block_edits: anchors are verified, bad edits are refused, overlapping
    ranges are dropped (keep the first). Cell writes don't shift other cells,
    so — unlike line edits — surviving edits can be applied in any order."""
    grid = _parse_grid(content)
    rows = len(grid)
    cols = len(grid[0]) if grid else 0

    normalized, rejected = [], []
    for e in edits:
        range_str = e.get("range", "")
        rng = _parse_range(range_str)
        if not rng:
            rejected.append({"edit": e, "reason": f"invalid range {range_str!r}"})
            continue
        r0, c0, r1, c1 = rng

        anchor = grid[r0][c0] if r0 < rows and c0 < cols else ""
        expect = e.get("expect")
        if expect and _norm(expect) not in _norm(anchor) and _norm(anchor) not in _norm(expect):
            rejected.append({
                "edit": e,
                "reason": f"anchor mismatch at {range_str}: expected {expect!r}, found {anchor!r}",
            })
            continue

        replacement = e.get("replacement", "")
        if replacement is None:
            replacement = ""
        want_rows, want_cols = r1 - r0 + 1, c1 - c0 + 1
        if replacement == "":
            rep_grid = [["" for _ in range(want_cols)] for _ in range(want_rows)]
        else:
            rep_grid = [r.split(",") for r in replacement.split("\n")]
            if len(rep_grid) != want_rows or any(len(rr) != want_cols for rr in rep_grid):
                rejected.append({
                    "edit": e,
                    "reason": f"replacement shape does not match range {range_str} "
                              f"({want_rows}x{want_cols} expected)",
                })
                continue

        normalized.append({
            "r0": r0, "c0": c0, "r1": r1, "c1": c1,
            "rep": rep_grid, "range": range_str, "cleared": replacement == "",
        })

    # Overlapping ranges would clobber each other — keep the first, drop the rest.
    kept, occupied = [], set()
    for ed in normalized:
        cells = {(r, c) for r in range(ed["r0"], ed["r1"] + 1) for c in range(ed["c0"], ed["c1"] + 1)}
        if cells & occupied:
            rejected.append({"edit": ed, "reason": f"range {ed['range']} overlaps an earlier edit"})
            continue
        occupied |= cells
        kept.append(ed)

    # Grow the grid if any surviving edit reaches beyond its current bounds.
    max_r = max((ed["r1"] for ed in kept), default=-1)
    max_c = max((ed["c1"] for ed in kept), default=-1)
    if max_r >= rows:
        grid.extend([[""] * cols for _ in range(max_r - rows + 1)])
        rows = len(grid)
    if max_c >= cols:
        for r in grid:
            r.extend([""] * (max_c - cols + 1))
        cols = max_c + 1

    applied = []
    for ed in kept:
        for i in range(ed["r1"] - ed["r0"] + 1):
            for j in range(ed["c1"] - ed["c0"] + 1):
                grid[ed["r0"] + i][ed["c0"] + j] = ed["rep"][i][j]
        applied.append({"range": ed["range"], "cleared": ed["cleared"]})

    return _grid_to_csv(grid), applied, rejected


@router.post("/canvas/{conversation_id}/{artifact_id}/patch")
async def patch_canvas(conversation_id: str, artifact_id: str, body: CanvasPatch):
    """Edit the canvas by line-addressed patch — only the named lines change."""
    from langchain_core.messages import SystemMessage, HumanMessage
    from app.llm import get_llm
    from app.storage import load_prefs

    convos = load_convos()
    _, art = _locate(convos, conversation_id, artifact_id)
    content = art.get("content") or ""
    instruction = (body.instruction or "").strip()
    if not instruction:
        raise HTTPException(status_code=400, detail="instruction is required")

    mode = (art.get("mode") or "code").lower()
    is_doc = mode == "document" and _looks_like_html(content)
    is_sheet = mode == "spreadsheet"

    if is_doc:
        # Prose is addressed by block; see the note above split_html_blocks.
        numbered = _number_blocks(content)
        if not numbered.strip():
            raise HTTPException(status_code=422, detail="Document has no editable blocks")
        system = (
            "You edit a document by emitting a minimal PATCH, never by rewriting it.\n"
            "You are given the document as numbered BLOCKS (paragraphs, headings, "
            "lists). Return ONLY a JSON object, no prose and no code fences:\n"
            '{"edits":[{"block":2,"expect":"<start of that block\'s text>",'
            '"replacement":"the new text for that block"}]}\n'
            "Rules:\n"
            "- Edit as FEW blocks as possible. Never restate unchanged blocks.\n"
            "- 'block' is the 1-based number shown in brackets.\n"
            "- 'expect' MUST match the start of that block's existing text. It is "
            "verified; a wrong anchor causes the edit to be rejected.\n"
            "- 'replacement' is the block's new PLAIN TEXT. Inline **bold**, "
            "*italic* and `code` are allowed. Do NOT include HTML tags — the "
            "block keeps its existing type (paragraph stays a paragraph, heading "
            "stays a heading).\n"
            "- For a list block, put one item per line.\n"
            '- Use "" as the replacement to delete a block.'
        )
        human = f"INSTRUCTION: {instruction}\n\nDOCUMENT:\n{numbered}"
        raw_kind = "block"
    elif is_sheet:
        grid = _parse_grid(content)
        numbered = _numbered_grid(grid)
        focus = ""
        if body.range:
            focus = f"\nThe user has range {body.range} selected. Prefer edits there.\n"
        system = (
            "You edit a spreadsheet by emitting a minimal PATCH, never by rewriting it.\n"
            "You are given the grid with column letters and row numbers (these "
            "coordinates are NOT part of any cell's value). Return ONLY a JSON "
            "object, no prose and no code fences:\n"
            '{"edits":[{"range":"B2","expect":"<exact current text of B2>",'
            '"replacement":"new value"},'
            '{"range":"C3:C5","expect":"<exact current text of C3>",'
            '"replacement":"10\\n20\\n30"}]}\n'
            "Rules:\n"
            "- 'range' is a single cell (\"B2\") or a rectangular range (\"B2:D5\") "
            "in A1 notation. Columns are letters, rows are 1-based numbers.\n"
            "- 'expect' MUST be the exact current text of the range's TOP-LEFT "
            "cell. It is verified; a wrong anchor causes the edit to be rejected.\n"
            "- For a single cell, 'replacement' is just the new value.\n"
            "- For a range, 'replacement' has one line per row and cells within a "
            "row separated by commas — it MUST match the range's dimensions "
            "exactly. Cell values themselves must not contain commas or newlines.\n"
            "- A formula starts with \"=\", e.g. \"=SUM(A1:A10)\", \"=B2*C2\".\n"
            "- Use \"\" as the replacement to clear a cell or range.\n"
            "- Change as FEW cells as possible. Never restate unchanged cells.\n"
            "- Emit multiple small edits rather than one giant range."
        )
        human = f"INSTRUCTION: {instruction}\n{focus}\nSPREADSHEET:\n{numbered}"
        raw_kind = "cell"
    else:
        numbered = _number_lines(content)
        focus = ""
        if body.start_line:
            focus = (f"\nThe user is focused on lines {body.start_line}"
                     f"-{body.end_line or body.start_line}. Prefer edits there.\n")
        raw_kind = "line"
        system = (
            "You edit files by emitting a minimal PATCH, never by rewriting the file.\n"
            "You are given the file with line numbers (the numbers are NOT part of the "
            "content). Return ONLY a JSON object, no prose and no code fences:\n"
            '{"edits":[{"start_line":12,"end_line":15,'
            '"expect":"<exact text of line 12>","replacement":"new line\\nanother line"}]}\n'
            "Rules:\n"
            "- Change as FEW lines as possible. Never restate unchanged lines.\n"
            "- start_line/end_line are inclusive and 1-based, over the ORIGINAL numbering.\n"
            "- 'expect' MUST be the exact original text of start_line. It is verified; "
            "a wrong anchor causes the edit to be rejected.\n"
            "- 'replacement' is the new text WITHOUT line numbers. Use \\n for multiple "
            "lines. Use \"\" to delete the range.\n"
            "- To INSERT without deleting, set end_line = start_line - 1; the text is "
            "inserted before start_line.\n"
            "- Preserve the file's existing indentation style exactly.\n"
            "- Emit multiple small edits rather than one giant range."
        )
        human = f"INSTRUCTION: {instruction}\n{focus}\nFILE:\n{numbered}"

    try:
        llm = get_llm(load_prefs())
        resp = await llm.ainvoke([SystemMessage(content=system), HumanMessage(content=human)])
        raw = _strip_wrapper(getattr(resp, "content", "") or "")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Patch generation failed: {e}")

    # Models sometimes wrap the JSON in a sentence; pull out the object.
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", raw, re.S)
        if not m:
            raise HTTPException(status_code=502, detail=f"Model did not return JSON: {raw[:200]}")
        try:
            data = json.loads(m.group(0))
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=502, detail=f"Malformed patch JSON: {e}")

    edits = data.get("edits") or []
    if not isinstance(edits, list) or not edits:
        raise HTTPException(status_code=502, detail="Model returned no edits")

    if is_doc:
        new_content, applied, rejected = apply_block_edits(content, edits)
    elif is_sheet:
        new_content, applied, rejected = apply_cell_edits(content, edits)
    else:
        new_content, applied, rejected = apply_line_edits(content, edits)

    if not applied:
        detail = rejected[0]["reason"] if rejected else "no applicable edits"
        raise HTTPException(status_code=422, detail=f"Patch could not be applied: {detail}")

    # Persist immediately so the stored copy and the editor agree.
    art["content"] = new_content
    art["edited_by_user"] = True
    save_convos(convos)

    return {
        "content": new_content,
        "applied": applied,
        "rejected": [r["reason"] for r in rejected],
        # Tells the client how to apply it: line ranges (Monaco), whole HTML
        # (Quill, which has no line model to splice into), or a full grid
        # reload (x-spreadsheet, same reasoning as Quill).
        "kind": raw_kind,
    }


# ── Export ────────────────────────────────────────────────────────────

def _build_docx(blocks, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()
    for kind, text, level in blocks:
        text = _strip_inline_md(text)
        if kind == "heading":
            doc.add_heading(text, level=min(max(level, 1), 6))
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        elif kind == "quote":
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(blocks, title: str) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    mono = ParagraphStyle("mono", parent=styles["Code"], fontName="Courier", fontSize=8.5, leading=11)
    quote = ParagraphStyle("quote", parent=styles["BodyText"], leftIndent=18,
                           textColor="#555555", fontName="Helvetica-Oblique")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER, title=title,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
    )

    flow = []
    for kind, text, level in blocks:
        text = _strip_inline_md(text)
        if kind == "code":
            # Preformatted keeps whitespace; it must not be XML-escaped twice.
            flow.append(Preformatted(text or " ", mono))
            continue
        safe = escape(text)
        if kind == "heading":
            flow.append(Spacer(1, 8))
            flow.append(Paragraph(safe, styles[f"Heading{min(max(level, 1), 4)}"]))
        elif kind in ("bullet", "number"):
            flow.append(Paragraph(f"• {safe}" if kind == "bullet" else safe,
                                  styles["BodyText"], bulletText=None))
        elif kind == "quote":
            flow.append(Paragraph(safe, quote))
        else:
            flow.append(Paragraph(safe, styles["BodyText"]))

    doc.build(flow or [Paragraph("", styles["BodyText"])])
    return buf.getvalue()


def _blocks_to_markdown(blocks) -> str:
    out = []
    for kind, text, level in blocks:
        if kind == "heading":
            out.append(f"{'#' * min(max(level, 1), 6)} {text}")
        elif kind == "bullet":
            out.append(f"- {text}")
        elif kind == "number":
            out.append(f"1. {text}")
        elif kind == "quote":
            out.append(f"> {text}")
        elif kind == "code":
            out.append(text)
        else:
            out.append(text)
        out.append("")
    return "\n".join(out).strip() + "\n"


def _coerce_cell(v: str):
    """CSV cells are always strings; write numbers back as numbers so
    formulas referencing them compute correctly, and pass formulas through
    as-is (a string starting with "=" is exactly how openpyxl recognises one)."""
    if v is None or v == "" or v.startswith("="):
        return v or None
    try:
        return int(v)
    except ValueError:
        pass
    try:
        return float(v)
    except ValueError:
        return v


def _build_xlsx(grid: list, title: str) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = (re.sub(r"[\[\]:*?/\\]", "_", title)[:31] or "Sheet1")
    for row in grid:
        ws.append([_coerce_cell(c) for c in row])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _grid_to_html_table(grid: list, title: str) -> str:
    rows_html = "".join(
        "<tr>" + "".join(f"<td>{html_escape(c)}</td>" for c in row) + "</tr>"
        for row in grid
    )
    return (
        f"<!doctype html><meta charset=\"utf-8\"><title>{html_escape(title)}</title>"
        f"<style>table{{border-collapse:collapse}}td{{border:1px solid #ccc;"
        f"padding:4px 8px;font:13px sans-serif}}</style>"
        f"<table>{rows_html}</table>"
    )


MEDIA_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "html": "text/html; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv; charset=utf-8",
}

# Formats each mode can actually produce — xlsx/csv only make sense for a
# grid, and the block/blocks writers below don't know how to lay out a table.
_SHEET_FORMATS = {"xlsx", "csv", "html"}
_TEXT_FORMATS = {"docx", "pdf", "md", "html", "txt"}


@router.get("/canvas/{conversation_id}/{artifact_id}/export")
async def export_canvas(conversation_id: str, artifact_id: str, format: str = "docx"):
    """Render a canvas to a downloadable file. DOCX is the default (xlsx for spreadsheets)."""
    fmt = (format or "docx").lower()
    if fmt not in MEDIA_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")

    _, art = _locate(load_convos(), conversation_id, artifact_id)
    content = art.get("content") or ""
    mode = (art.get("mode") or "code").lower()
    title = art.get("title") or "canvas"
    safe_name = re.sub(r"[^\w.-]+", "_", title).strip("_") or "canvas"
    is_sheet = mode == "spreadsheet"

    if is_sheet and fmt not in _SHEET_FORMATS:
        raise HTTPException(status_code=400, detail=f"{fmt} is not available for spreadsheets")
    if not is_sheet and fmt not in _TEXT_FORMATS:
        raise HTTPException(status_code=400, detail=f"{fmt} is only available for spreadsheets")

    if is_sheet:
        grid = _parse_grid(content)
        if fmt == "csv":
            data = content.encode()
        elif fmt == "html":
            data = _grid_to_html_table(grid, title).encode()
        else:  # xlsx
            try:
                data = _build_xlsx(grid, title)
            except ImportError as e:
                raise HTTPException(status_code=500, detail=f"Exporter unavailable: {e}")
    elif fmt == "html":
        body = content if _looks_like_html(content) else f"<pre>{content}</pre>"
        data = (
            f"<!doctype html><meta charset=\"utf-8\"><title>{title}</title>"
            f"<body>{body}</body>"
        ).encode()
    elif fmt == "txt":
        blocks = _to_blocks(content, mode)
        data = "\n".join(b[1] for b in blocks).encode()
    elif fmt == "md":
        data = (content if mode == "code" else
                _blocks_to_markdown(_to_blocks(content, mode))).encode()
    else:
        blocks = _to_blocks(content, mode)
        try:
            data = _build_docx(blocks, title) if fmt == "docx" else _build_pdf(blocks, title)
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Exporter unavailable: {e}")

    return Response(
        content=data,
        media_type=MEDIA_TYPES[fmt],
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.{fmt}"'},
    )
